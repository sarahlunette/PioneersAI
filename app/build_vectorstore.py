"""
Build Qdrant Vectorstore (v4)
- pdfplumber for PDF extraction
- Proper ingestion into Qdrant (manual chunk loop)
- Supports TXT, MD, DOCX, PDF, CSV rows
- LOCAL embedding model: ../models/all-MiniLM-L6-v2
"""

from pathlib import Path
import pandas as pd
import pdfplumber

from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance, PointStruct

from llama_index.core import Document
from llama_index.embeddings.huggingface import HuggingFaceEmbedding


# ============================================================
# CONFIG
# ============================================================
DOCS_DIR = Path("./docs")

QDRANT_URL = "http://localhost:6333"
COLLECTION_NAME = "island_docs"

EMBEDDING_MODEL = "../models/all-MiniLM-L6-v2"


# ============================================================
# PDF extraction using pdfplumber
# ============================================================
def load_pdf(path: Path) -> list[Document]:
    docs = []
    try:
        with pdfplumber.open(path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""

        if text.strip():
            docs.append(Document(text=text, metadata={"source": path.name}))
        else:
            print(f"⚠ PDF EMPTY (no text extracted): {path.name}")

    except Exception as e:
        print(f"❌ Error reading PDF {path.name}: {e}")

    return docs


# ============================================================
# CSV ingestion
# ============================================================
def load_csv_rows(path: Path):
    df = pd.read_csv(path)
    docs = []

    for i, row in df.iterrows():
        text = "\n".join(f"{col}: {row[col]}" for col in df.columns)
        if text.strip():
            docs.append(
                Document(
                    text=text,
                    metadata={"source": path.name, "row_index": int(i)}
                )
            )
    return docs


# ============================================================
# TXT / MD / DOCX (simple text load)
# ============================================================
def load_text_file(path: Path):
    try:
        text = path.read_text(errors="ignore")
        if text.strip():
            return [Document(text=text, metadata={"source": path.name})]
    except:
        pass
    return []


# ============================================================
# MAIN
# ============================================================
def main():
    print("\n=== 🚀 QDRANT VECTORSTORE BUILDER v4 ===")

    if not DOCS_DIR.exists():
        raise FileNotFoundError("docs/ folder missing")

    all_docs = []

    print("📚 Loading documents…")
    for file in DOCS_DIR.iterdir():
        if file.suffix.lower() == ".pdf":
            print(f"  → PDF: {file.name}")
            all_docs.extend(load_pdf(file))

        elif file.suffix.lower() in [".txt", ".md"]:
            print(f"  → TEXT: {file.name}")
            all_docs.extend(load_text_file(file))

        elif file.suffix.lower() == ".docx":
            print(f"  → DOCX: {file.name}")
            # optional: python-docx extraction
            try:
                import docx
                doc = docx.Document(str(file))
                text = "\n".join([p.text for p in doc.paragraphs])
                all_docs.append(Document(text=text, metadata={"source": file.name}))
            except:
                print(f"❌ Cannot read DOCX: {file.name}")

        elif file.suffix.lower() == ".csv":
            print(f"  → CSV rows: {file.name}")
            all_docs.extend(load_csv_rows(file))

    print(f"\n📦 Total loaded documents: {len(all_docs)}")

    if not all_docs:
        print("⚠ No documents loaded. STOP.")
        return

    # ============================================================
    # Embedding model
    # ============================================================
    print(f"\n🧠 Loading embedding model: {EMBEDDING_MODEL}")
    embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)

    embedding_dim = len(embed_model.get_text_embedding("test"))
    print(f"🔢 Embedding size = {embedding_dim}")

    # ============================================================
    # Qdrant setup
    # ============================================================
    print("\n🗄 Connecting to Qdrant…")
    client = QdrantClient(url=QDRANT_URL)

    print(f"📌 Recreating collection `{COLLECTION_NAME}`…")
    client.recreate_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=embedding_dim, distance=Distance.COSINE)
    )

    # ============================================================
    # MANUAL INGESTION INTO QDRANT (IMPORTANT)
    # ============================================================
    print("\n⚙️ Inserting vectors into Qdrant…")

    points_to_insert = []
    point_id = 1

    for doc in all_docs:
        text = doc.text.strip()
        if not text:
            continue

        vector = embed_model.get_text_embedding(text)

        points_to_insert.append(
            PointStruct(
                id=point_id,
                vector=vector,
                payload={"text": text, **doc.metadata}
            )
        )
        point_id += 1

    if points_to_insert:
        client.upsert(collection_name=COLLECTION_NAME, points=points_to_insert)

    # ============================================================
    # Result
    # ============================================================
    count = client.count(COLLECTION_NAME).count

    print("\n🎉 DONE — Qdrant Vectorstore BUILT!")
    print(f"📌 Collection: {COLLECTION_NAME}")
    print(f"📊 Total vectors stored: {count}")
    print("🚀 Ready for Claude RAG.\n")


if __name__ == "__main__":
    main()
